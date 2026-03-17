"""
Generate summary of given JSON or given JSON Document in to Microsoft Word Documents.
This functionality uses a python module called

python-docx - a Python library for creating and updating Microsoft Word (.docx) files.

https://python-docx.readthedocs.io/

"""

from datetime import datetime
import os
from typing import Union

import docx
from docx import oxml
from docx import shared
from docx.enum import text
from docx.oxml import ns

from sarif import charts, sarif_file
from sarif.sarif_file_utils import combine_record_code_and_description


def generate_word_docs_from_sarif_inputs(
    input_files: sarif_file.SarifFileSet,
    image_file: Union[str, None],
    output: str,
    output_multiple_files: bool,
    date_val: datetime = datetime.now(),
):
    """
    Convert SARIF input to Word file output.
    """
    if not input_files:
        raise ValueError("No input files specified!")

    output_file = output
    output_file_name = output
    if output_multiple_files:
        for input_file in input_files:
            output_file_name = input_file.get_file_name_without_extension() + ".docx"
            print(
                "Writing Word summary of",
                input_file.get_file_name(),
                "to",
                output_file_name,
            )
            report = input_file.get_report()
            _generate_word_summary(
                input_file,
                report,
                os.path.join(output, output_file_name),
                image_file,
                date_val,
            )
        output_file_name = "static_analysis_output.docx"
        output_file = os.path.join(output, output_file_name)

    source_description = input_files.get_description()
    print("Writing Word summary of", source_description, "to", output_file_name)
    report = input_files.get_report()
    _generate_word_summary(input_files, report, output_file, image_file, date_val)


def _generate_word_summary(
    sarif_data, report, output_file, image_file: Union[str, None], date_val: datetime
):
    # Create a new document
    document = docx.Document()

    severities = report.get_severities()
    _add_heading_and_highlevel_info(
        document, sarif_data, report, severities, output_file, image_file, date_val
    )
    _dump_errors_summary_by_sev(document, report, severities)
    _dump_each_error_in_detail(document, report, severities)

    # finally, save the document.
    document.save(output_file)


def _add_heading_and_highlevel_info(
    document,
    sarif_data,
    report,
    severities,
    output_file,
    image_path: Union[str, None],
    date_val: datetime,
):
    tool_name = ", ".join(sarif_data.get_distinct_tool_names())
    heading = f"Sarif Summary: {tool_name}"

    if image_path:
        document.add_picture(image_path)
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading(heading, 0)
    document.add_paragraph(f"Document generated on: {date_val}")

    sevs = ", ".join(severities)
    document.add_paragraph(
        f"Total number of various severities ({sevs}): {sarif_data.get_result_count()}"
    )
    filter_stats = sarif_data.get_filter_stats()
    if filter_stats:
        document.add_paragraph(f"Results were filtered by {filter_stats}.")

    pie_chart_image_file_path = output_file.replace(".docx", "_severity_pie_chart.png")
    if charts.generate_severity_pie_chart(report, pie_chart_image_file_path):
        document.add_picture(pie_chart_image_file_path)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()


def _dump_errors_summary_by_sev(document, report, severities):
    """
    For each severity level (in priority order): create a list of the errors of
    that severity, print out how many there are and then do some further analysis
    of which error codes are present.
    """
    for severity in severities:
        errors_of_severity = report.get_issue_type_count_for_severity(severity)
        document.add_heading(f"Severity : {severity} [ {errors_of_severity} ]", level=1)
        sorted_dict = report.get_issue_type_histogram_for_severity(severity)
        if sorted_dict:
            for key, count in sorted_dict.items():
                document.add_paragraph(f"{key}: {count}", style="List Bullet")
        else:
            document.add_paragraph("None", style="List Bullet")


def _dump_each_error_in_detail(document, report, severities):
    """
    Write out the errors to a table so that a human can do further analysis.
    """
    document.add_page_break()

    for severity in severities:
        errors_of_severity = report.get_issues_for_severity(severity)
        # Sample:
        # [{'Location': 'C:\\Max\\AccessionAndroid\\scripts\\parse_coverage.py', 'Line': 119,
        #       'Severity': 'error', 'Code': 'DS126186 Disabled certificate validation'},
        # {'Location': 'C:\\Max\\AccessionAndroid\\scripts\\parse_code_stats.py', 'Line': 61,
        #       'Severity': 'error', 'Code': 'DS126186 Disabled certificate validation'},
        # ]
        if errors_of_severity:
            document.add_heading(f"Severity : {severity}", level=2)
            table = document.add_table(rows=1 + len(errors_of_severity), cols=3)

            table.style = "Table Grid"  # ColorfulGrid-Accent5'
            table.autofit = False

            table.alignment = text.WD_TAB_ALIGNMENT.CENTER

            # Cell widths
            widths = [shared.Inches(2), shared.Inches(4), shared.Inches(0.5)]

            # To avoid performance problems with large tables, prepare the entries first in this
            # list, then iterate the table cells and copy them in.
            # First populate the header row
            cells_text = ["Code", "Location", "Line"]

            hdr_cells = table.rows[0].cells
            for i in range(3):
                table.rows[0].cells[i]._tc.get_or_add_tcPr().append(
                    oxml.parse_xml(
                        r'<w:shd {} w:fill="5fe3d8"/>'.format(ns.nsdecls("w"))
                    )
                )
                run = hdr_cells[i].paragraphs[0].add_run(cells_text[i])
                run.bold = True
                hdr_cells[i].paragraphs[
                    0
                ].alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
                hdr_cells[i].width = widths[i]

            for eachrow in errors_of_severity:
                cells_text += [
                    combine_record_code_and_description(eachrow),
                    eachrow["Location"],
                    str(eachrow["Line"]),
                ]

            # Note: using private property table._cells to avoid performance issue.  See
            # https://stackoverflow.com/a/69105798/316578
            col_index = 0
            for cell, cell_text in zip(table._cells, cells_text):
                cell.text = cell_text
                cell.width = widths[col_index]
                col_index = col_index + 1 if col_index < 2 else 0
        else:
            document.add_heading(f"Severity : {severity}", level=2)
            document.add_paragraph("None", style="List Bullet")
